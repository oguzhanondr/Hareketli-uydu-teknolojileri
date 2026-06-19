from pathlib import Path
from datetime import date
from math import floor

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SCREENSHOTS = ROOT / "artifacts" / "ktr" / "screenshots"
WORKING = ROOT / "artifacts" / "ktr" / "working"
OUT_DOCX = ROOT / "output" / "docx" / "ARES-Reflect_Teknik_Sistem_Aciklama_Dosyasi.docx"

NAVY = "132033"
NAVY_2 = "1B2D46"
CYAN = "00B8D9"
BLUE = "2F6FFF"
PURPLE = "A855F7"
GREEN = "16A36A"
AMBER = "D9A300"
RED = "D94343"
INK = "1F2937"
MUTED = "5D6878"
LIGHT = "F3F6FA"
LIGHT_BLUE = "EAF2FF"
LIGHT_CYAN = "E9F9FC"
LIGHT_AMBER = "FFF7DF"
WHITE = "FFFFFF"
GRID = "CAD3DF"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=GRID, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names += ["insideH", "insideV"]
    for edge in names:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_fixed_width(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if table.rows:
        set_repeat_table_header(table.rows[0])


def set_font(run, name="Arial", size=10.5, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color=CYAN, size=10, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=MUTED)


def add_run_text(paragraph, text, bold=False, italic=False, color=INK, size=10.5):
    r = paragraph.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return r


def add_body(doc, text, bold_lead=None, after=7, keep=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        add_run_text(p, bold_lead, bold=True)
        add_run_text(p, text[len(bold_lead):])
    else:
        add_run_text(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_run_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_run_text(p, text)
    return p


def add_heading(doc, text, level=1, number=None):
    label = f"{number} {text}" if number else text
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(label)
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: NAVY, 2: BLUE, 3: NAVY_2}
    set_font(r, size=sizes[level], color=colors[level], bold=True)
    return p


def add_caption(doc, text, source=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3 if source else 8)
    p.paragraph_format.keep_with_next = bool(source)
    add_run_text(p, text, bold=True, size=9, color=NAVY)
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        add_run_text(p2, source, italic=True, size=8, color=MUTED)


def add_figure(doc, image_path, caption, width=6.35, source="Kaynak: ARES-Reflect uygulamasından 19 Haziran 2026 tarihinde alınan ekran görüntüsü."):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    shape = r.add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    add_caption(doc, caption, source)


def add_callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    set_table_borders(table, color=accent, size=8, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run_text(p, title, bold=True, color=accent, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run_text(p2, text, size=9.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows, widths=(2500, 6860), header=None):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    if header:
        cells = table.rows[0].cells
        cells[0].merge(cells[1])
        set_cell_shading(cells[0], NAVY)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_text(p, header, bold=True, color=WHITE, size=10)
        set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT)
        p = cells[0].paragraphs[0]
        add_run_text(p, label, bold=True, color=NAVY, size=9)
        p2 = cells[1].paragraphs[0]
        add_run_text(p2, value, size=9)
    set_table_fixed_width(table, list(widths))
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix_table(doc, headers, rows, widths, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_text(p, header, bold=True, color=WHITE, size=font_size)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], "F8FAFC")
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx else WD_ALIGN_PARAGRAPH.CENTER
            add_run_text(p, str(value), size=font_size)
    set_table_fixed_width(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def make_assets():
    WORKING.mkdir(parents=True, exist_ok=True)

    # Bölüm 7 figürleri, canlı dağıtımdan (hareketli-uydu-teknolojileri.vercel.app)
    # alınan tek bir gerçek ekran görüntüsünden türetilir; böylece harita, sonuç
    # paneli ve seçilen IRS detayı aynı senaryoya ait ve birbiriyle tutarlı kalır.
    real_source = WORKING / "07-gercek-uygulama-sonuc.png"
    if real_source.exists():
        real = Image.open(real_source)
        real.crop((0, 66, 1919, 1078)).save(WORKING / "07a-sonuc-tam.png")
        real.crop((1163, 200, 1919, 862)).save(WORKING / "07b-terminal-irs-paneli.png")
        real.crop((1163, 868, 1919, 1030)).save(WORKING / "07c-irs-detay.png")

    width, height = 1800, 950
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_regular = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 31)
    font_small = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 24)
    font_bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 34)
    font_title = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 44)
    colors = {
        "navy": "#132033",
        "blue": "#2f6fff",
        "purple": "#a855f7",
        "cyan": "#00a8c6",
        "green": "#16a36a",
        "light": "#eef3f8",
        "line": "#b8c4d2",
        "text": "#1f2937",
    }

    draw.text((70, 45), "ARES-Reflect Sistem Mimarisi", font=font_title, fill=colors["navy"])
    draw.text(
        (70, 105),
        "Fiziksel terminal vizyonu ile deterministik saha yerleşim yazılımının ilişkisi",
        font=font_regular,
        fill="#536174",
    )

    boxes = [
        (80, 230, 390, 430, "Saha Girdileri", "Depremzede noktaları\nEnkaz işaretleri\nBina ayak izleri", colors["cyan"]),
        (475, 230, 790, 430, "Yerel Geometri Motoru", "Deterministik k-means\nTerminal aday taraması\n2B LoS ve koridor testi", colors["blue"]),
        (875, 230, 1190, 430, "Ortak Optimizasyon", "Terminal + IRS birlikte\nMutlak kalite puanı\nMinimum yeterli IRS seti", colors["purple"]),
        (1275, 230, 1590, 430, "Operasyonel Çıktı", "Harita yerleşimi\nMetrik ve gerekçeler\nPDF saha raporu", colors["green"]),
        (475, 610, 790, 800, "Fiziksel Terminal", "BLDC + enkoder\nIMU + EKF\nKaskad stabilizasyon", colors["navy"]),
        (875, 610, 1190, 800, "RIS/IRS Katmanı", "İki bacaklı koridor\nCephe/direk montajı\nElektronik yönlendirme", colors["purple"]),
        (1275, 610, 1590, 800, "Gemini Destek Katmanı", "Açıklama\nGeçerli sonuçları sıralama\nİkincil görsel denetim", colors["cyan"]),
    ]
    for x1, y1, x2, y2, title, body, accent in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=colors["light"], outline=accent, width=5)
        draw.rectangle((x1, y1, x2, y1 + 18), fill=accent)
        draw.text((x1 + 24, y1 + 42), title, font=font_bold, fill=colors["navy"])
        for idx, line in enumerate(body.split("\n")):
            draw.text((x1 + 24, y1 + 98 + idx * 42), line, font=font_small, fill=colors["text"])

    def arrow(a, b, color="#65758b"):
        draw.line((a, b), fill=color, width=7)
        ax, ay = b
        draw.polygon([(ax, ay), (ax - 20, ay - 13), (ax - 20, ay + 13)], fill=color)

    arrow((390, 330), (475, 330))
    arrow((790, 330), (875, 330))
    arrow((1190, 330), (1275, 330))
    arrow((630, 430), (630, 610))
    arrow((1030, 430), (1030, 610))
    arrow((1430, 610), (1430, 445), color=colors["cyan"])
    draw.text(
        (70, 875),
        "Not: Gemini koordinat üretmez; yerleşim ve fiziksel geçerlilik her zaman yerel motor tarafından belirlenir.",
        font=font_regular,
        fill=colors["navy"],
    )
    img.save(WORKING / "sistem-mimarisi.png")

    width, height = 1800, 840
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 40), "Deterministik Yerleşim İşlem Hattı", font=font_title, fill=colors["navy"])
    steps = [
        ("1", "Sahne modeli", "Depremzede, enkaz\nve bina verisi"),
        ("2", "Kümeleme", "Üç deterministik\nhedef bölgesi"),
        ("3", "Terminal taraması", "Izgara, halka ve\nkoridor adayları"),
        ("4", "IRS üretimi", "Yön, mesafe ve\ncephe adayları"),
        ("5", "Fiziksel eleme", "İki bacak LoS,\nmesafe, kalite"),
        ("6", "Ortak seçim", "Kapsama tavanına\nulaşan en küçük set"),
        ("7", "Sunum", "Harita, kartlar,\nmetrikler ve PDF"),
    ]
    start_x, gap, box_w, box_h, y = 55, 28, 220, 380, 210
    for idx, (num, title, body) in enumerate(steps):
        x = start_x + idx * (box_w + gap)
        accent = [colors["cyan"], colors["blue"], colors["blue"], colors["purple"], "#d94343", colors["green"], colors["navy"]][idx]
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=25, fill=colors["light"], outline=accent, width=5)
        draw.ellipse((x + 72, y + 25, x + 148, y + 101), fill=accent)
        bbox = draw.textbbox((0, 0), num, font=font_bold)
        draw.text((x + 110 - (bbox[2] - bbox[0]) / 2, y + 42), num, font=font_bold, fill="white")
        title_lines = title.split(" ")
        if len(title) > 17 and len(title_lines) > 1:
            midpoint = len(title_lines) // 2
            title_text = " ".join(title_lines[:midpoint]) + "\n" + " ".join(title_lines[midpoint:])
        else:
            title_text = title
        draw.multiline_text((x + 18, y + 125), title_text, font=font_bold, fill=colors["navy"], spacing=5, align="center")
        draw.multiline_text((x + 20, y + 245), body, font=font_small, fill=colors["text"], spacing=12, align="center")
        if idx < len(steps) - 1:
            arrow((x + box_w, y + box_h // 2), (x + box_w + gap - 5, y + box_h // 2))
    draw.text(
        (70, 700),
        "Geçersiz adaylar sunum katmanına ulaşmadan elenir. Aynı girdi kümesi aynı yerel yerleşimi üretir.",
        font=font_regular,
        fill=colors["navy"],
    )
    img.save(WORKING / "islem-hatti.png")


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    style_values = {
        "Heading 1": (16, NAVY, 15, 7),
        "Heading 2": (13, BLUE, 11, 5),
        "Heading 3": (11.5, NAVY_2, 8, 4),
    }
    for name, (size, color, before, after) in style_values.items():
        s = styles[name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = rgb(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        s = styles[list_style]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        s.font.size = Pt(10.5)
        s.paragraph_format.space_after = Pt(4)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run_text(p, "ARES-Reflect", bold=True, color=NAVY, size=9)
    add_run_text(p, "  |  Teknik Sistem Açıklama Dosyası", color=MUTED, size=8.5)
    set_paragraph_border(p, color=CYAN, size=6, space=3)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_fixed_width(table, [7000, 2360], indent=0)
    set_table_borders(table, color=WHITE, size=0, inside=False)
    p1 = table.cell(0, 0).paragraphs[0]
    add_run_text(p1, "KTR hazırlığı için teknik kaynak belge", size=8, color=MUTED)
    p2 = table.cell(0, 1).paragraphs[0]
    add_page_number(p2)
    return doc


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "TEKNİK SİSTEM AÇIKLAMA DOSYASI", bold=True, color=CYAN, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_run_text(p, "ARES-Reflect", bold=True, color=NAVY, size=30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    add_run_text(
        p,
        "Afet Sahalarında Hareketli Uydu Terminali ve IRS Yerleşim Karar Destek Sistemi",
        bold=True,
        color=NAVY_2,
        size=15,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    add_run_text(
        p,
        "Kritik Tasarım Raporu yazımında kullanılmak üzere hazırlanmış bağımsız teknik kaynak",
        italic=True,
        color=MUTED,
        size=10.5,
    )

    add_figure(
        doc,
        WORKING / "sistem-mimarisi.png",
        "Şekil 1. ARES-Reflect sistem mimarisinin kavramsal gösterimi",
        width=6.3,
        source="Kaynak: Proje yazılımı, ÖTR içeriği ve uygulama mimarisi esas alınarak hazırlanmıştır.",
    )

    table = doc.add_table(rows=4, cols=2)
    cover_rows = [
        ("Belgenin amacı", "Sistemin ne olduğunu, nasıl çalıştığını ve arayüz çıktılarının mühendislik anlamını açıklamak"),
        ("Hedef kullanıcı", "KTR belgesini hazırlayacak proje ekibi ve teknik değerlendiriciler"),
        ("İnceleme tarihi", "19 Haziran 2026"),
        ("Dayanak", "Güncel kaynak kod, çalışan uygulama, regresyon testleri ve ÖTR dokümanı"),
    ]
    for idx, (label, value) in enumerate(cover_rows):
        c0, c1 = table.rows[idx].cells
        set_cell_shading(c0, LIGHT)
        add_run_text(c0.paragraphs[0], label, bold=True, color=NAVY, size=9)
        add_run_text(c1.paragraphs[0], value, size=9)
    set_table_fixed_width(table, [2200, 7160])
    set_table_borders(table)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "Belgenin Kullanım Biçimi", 1)
    add_callout(
        doc,
        "Bu belge KTR şablonu değildir.",
        "Belge, proje ekibinin hazırlayacağı Kritik Tasarım Raporu için doğrulanmış teknik içerik, açıklama dili, ekran görüntüsü ve sınırlandırma kaynağıdır. Bölüm sırası yarışma şablonunu taklit etmez; sistemi anlaşılır ve izlenebilir biçimde açıklar.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    add_body(
        doc,
        "Metindeki “uygulama” ifadesi, afet sahasındaki terminal ve IRS yerleşimini hesaplayan ARES-Reflect web arayüzünü; “fiziksel sistem” ifadesi ise hareketli uydu terminali, stabilizasyon platformu, sensör füzyonu ve RIS/IRS donanım vizyonunu ifade eder. Bu iki katman ilişkili olmakla birlikte aynı şey değildir."
    )
    add_heading(doc, "İçerik Haritası", 2)
    items = [
        "1. Yönetici özeti ve sistemin kapsamı",
        "2. Operasyonel problem ve kullanım senaryosu",
        "3. Fiziksel terminal ile yerleşim yazılımının ilişkisi",
        "4. Kullanıcı arayüzü ve operatör iş akışı",
        "5. Deterministik yerleşim algoritması",
        "6. Terminal ve IRS kalite/uygunluk metrikleri",
        "7. Çalışan uygulamadan örnek senaryo incelemesi",
        "8. Gemini destek katmanının rolü ve yetki sınırı",
        "9. Doğrulama, testler, varsayımlar ve sınırlamalar",
        "10. KTR yazımında kullanılabilecek teknik ifadeler ve veri sözlüğü",
    ]
    for item in items:
        add_bullet(doc, item)
    add_heading(doc, "Kanıt ve Kaynak Hiyerarşisi", 2)
    add_number(doc, "Çalışan uygulama ve güncel kaynak kod, yazılım davranışı için birincil kanıttır.")
    add_number(doc, "19 Haziran 2026 tarihli regresyon çıktıları, belirlenmiş geometrik kuralların test kanıtıdır.")
    add_number(doc, "ÖTR dokümanı, fiziksel terminal vizyonu ve tasarım hedefleri için kaynak kabul edilmiştir.")
    add_number(doc, "Tam RF link bütçesi, 3B ışın izleme veya fiziksel prototip ölçümü bulunmayan alanlar açıkça model/varsayım olarak işaretlenmiştir.")
    doc.add_page_break()


def add_main_content(doc):
    add_heading(doc, "Yönetici Özeti", 1, "1.")
    add_body(
        doc,
        "ARES-Reflect, afet sonrasında karasal haberleşme altyapısının kesildiği ve doğrudan uydu görüşünün bina geometrileri nedeniyle zayıfladığı kentsel alanlarda, hareketli uydu terminali ile yeniden yapılandırılabilir akıllı yüzeylerin (RIS/IRS) kurulacağı noktaları belirleyen bir karar destek sistemidir. Uygulama; depremzede konumlarını, kullanıcı tarafından işaretlenen enkazları ve bina ayak izlerini birlikte değerlendirerek terminal–IRS–hedef biçimindeki iki bacaklı haberleşme koridorlarını üretir."
    )
    add_body(
        doc,
        "Sistemin temel tasarım ilkesi “önce fiziksel geçerlilik, sonra sıralama” yaklaşımıdır. Ayakta kalan bir bina terminal–IRS hattını kesiyorsa aday geçersiz kabul edilir. IRS–hedef tarafında açık erişilebilen en az bir depremzede bulunması, terminale asgari mesafe ve kalite eşiğinin sağlanması gerekir. Geçerli adaylar daha sonra kapsama, yol uzunluğu, yansıma verimi, cephe uyumu ve tahmini link kazancı gibi ölçütlerle puanlanır."
    )
    add_body(
        doc,
        "Yerleşim motoru deterministiktir: aynı depremzede, enkaz ve bina veri kümesi için rassal sayı kullanılmadan aynı yerel sonuç üretilir. Gemini entegrasyonu koordinat üretmez ve geçersiz bir adayı geçerli hale getiremez; yalnızca yerel motorun sonuçlarını açıklamak, zaten geçerli terminal kartlarının sunum sırasını iyileştirmek ve harita görüntüsüne ikincil görsel denetim uygulamak için kullanılır."
    )
    add_callout(
        doc,
        "Sistemin tek cümlelik tanımı",
        "ARES-Reflect, afet sahasındaki bina ve hedef geometrisini deterministik olarak analiz ederek uydu terminali ile minimum sayıda gerekli IRS biriminin fiziksel olarak uygulanabilir yerlerini öneren yerel-öncelikli bir karar destek yazılımıdır.",
        fill=LIGHT_CYAN,
        accent=CYAN,
    )

    add_heading(doc, "Sistemin Kapsamı ve Kapsam Dışı Hususlar", 2, "1.1")
    add_matrix_table(
        doc,
        ["Kapsam içi", "Kapsam dışı / ayrıca doğrulanmalı"],
        [
            ("2B bina ayak izlerine dayalı görüş hattı analizi", "Tam dalga elektromanyetik benzetim"),
            ("Terminal ve IRS aday koordinatı üretimi", "Uydu efemerisine dayalı gerçek zamanlı takip"),
            ("Cephe/direk montaj önerisi ve yaklaşık yükseklik", "Statik, ankraj ve saha güvenliği hesabı"),
            ("Sezgisel link kazancı ve mutlak kalite puanı", "Eksiksiz RF link bütçesi ve saha ölçümü"),
            ("Deterministik senaryo tekrarı ve regresyon testi", "Fiziksel prototip kabul testi"),
        ],
        [4680, 4680],
        font_size=9,
    )

    add_heading(doc, "Operasyonel Problem ve Kullanım Senaryosu", 1, "2.")
    add_body(
        doc,
        "Afet sonrasında baz istasyonları, enerji altyapısı ve karasal omurga bağlantıları hizmet dışı kalabilir. Aynı anda yoğun yapılaşma, enkaz ve ayakta kalan binalar uydu ile kullanıcı arasındaki doğrudan görüş hattını kısıtlayabilir. Projenin hedefi geniş bant internet sağlamak değil; SOS, GPS konumu ve kısa mesaj gibi düşük veri hızlı fakat kritik haberleşme paketleri için yeterli bağlantı koridoru oluşturmaktır."
    )
    add_body(
        doc,
        "ÖTR’de fiziksel terminalin hareketli kurtarma aracı üzerinde Satcom-on-the-Move yaklaşımıyla çalışması; iki eksenli doğrudan tahrikli BLDC platform, manyetik enkoder, dokuz eksenli IMU, Genişletilmiş Kalman Filtresi (EKF) ve kaskad kontrol ile yönelimin korunması öngörülmektedir. RIS/IRS katmanı ise doğrudan görüşün yetersiz olduğu bölgelerde RF sinyalini elektronik olarak yeniden yönlendiren pasif veya düşük güçlü bir röle görevi üstlenir."
    )
    add_kv_table(
        doc,
        [
            ("Görev verisi", "SOS, konum ve kısa mesaj sınıfı dar bant paketleri"),
            ("Terminal kütle sınırı", "20 kg (ÖTR tasarım kısıtı)"),
            ("Nominal güç sınırı", "140 W (ÖTR tasarım kısıtı)"),
            ("Hedef marjinal IRS kazancı", "5–10 dB; ÖTR’de örnek ortalama +7,5 dB"),
            ("Platform bozucusu", "±8° roll/pitch, 10 s periyotlu örnek bozucu"),
            ("Yeniden kilitlenme hedefi", "8 saniyenin altında"),
            ("Azimut hareketi", "Slipring ile 0–360° sürekli dönüş"),
        ],
        header="ÖTR’den Yazılıma Bağlanan Temel Tasarım Hedefleri",
    )
    add_callout(
        doc,
        "Kritik ayrım",
        "ÖTR’deki 5–10 dB hedefi fiziksel RF tasarım hedefidir. Uygulamadaki link_gain_db değeri ise yerleşim adaylarını karşılaştırmak için kullanılan kestirimsel bir metriktir; tek başına ölçülmüş RF kazancı olarak raporlanmamalıdır.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )

    add_heading(doc, "Fiziksel Terminal ile Yerleşim Yazılımının İlişkisi", 1, "3.")
    add_figure(
        doc,
        WORKING / "sistem-mimarisi.png",
        "Şekil 2. Fiziksel terminal, IRS ve yerleşim yazılımının katmanlı ilişkisi",
        source="Kaynak: Güncel yazılım mimarisi ve ÖTR içeriği esas alınarak hazırlanmıştır.",
    )
    add_heading(doc, "Fiziksel Alt Sistem", 2, "3.1")
    add_body(
        doc,
        "Fiziksel terminal; uyduya veya belirlenen IRS yönüne dönük kalması gereken hareketli bir anten platformudur. Doğrudan tahrikli BLDC motorların kullanılması, redüktör boşluğunu azaltarak yönelim hassasiyetini artırmayı hedefler. Enkoderler motor eksenlerini, IMU ise platform hareketini ölçer. EKF, ivmeölçer ve jiroskop verilerini birleştirerek gürültü ve sürüklenmeye daha dayanıklı bir yönelim kestirimi üretir."
    )
    add_body(
        doc,
        "Kaskad kontrol mimarisinde dış çevrim konum/yönelim hatasını, iç çevrim ise açısal hız ve hızlı bozucuları kontrol eder. Yerleşim yazılımı bu kontrol döngüsünün yerine geçmez; kontrol sistemine “terminal ve IRS geometrisi hangi doğrultuda kurulmalı?” sorusunun saha cevabını sağlar."
    )
    add_heading(doc, "Yerleşim Yazılımı", 2, "3.2")
    add_body(
        doc,
        "Yazılımın girdisi sensör telemetrisi değil, saha geometrisidir. Depremzede işaretleri, enkaz kimlikleri ve bina poligonları kullanılarak olası terminal/IRS noktaları üretilir. Çıktı; koordinat, kalite, kapsama, montaj yönü, yaklaşık montaj yüksekliği ve seçilme gerekçesidir. Bu çıktı fiziksel ekip için keşif ve kurulum kararı niteliğindedir."
    )

    add_heading(doc, "Kullanıcı Arayüzü ve Operatör İş Akışı", 1, "4.")
    add_figure(
        doc,
        SCREENSHOTS / "01-analiz-genel-gorunum.png",
        "Şekil 3. ARES-Reflect başlangıç ekranı ve ana çalışma alanları",
    )
    add_heading(doc, "Üst Durum Çubuğu", 2, "4.1")
    add_kv_table(
        doc,
        [
            ("Yerleşim: Yerel", "Terminal ve IRS koordinatlarının deterministik yerel geometri motorundan geldiğini gösterir."),
            ("Açıklama", "IRS açıklamasının yerel metin üreticisinden mi Gemini’den mi geldiğini gösterir."),
            ("Rerank", "Geçerli terminal kartlarının sunum sırasının yerel veya Gemini kaynaklı olduğunu belirtir."),
            ("Doğrulama", "Harita ekran görüntüsüne uygulanan ikincil görsel denetimin durumunu gösterir."),
            ("Proje Detayları", "ÖTR hedefleri, algoritma akışı ve veri sözlüğünü açıklayan bilgi sayfasını açar."),
            ("PDF Dışa Aktar", "Mevcut senaryonun terminal/IRS metriklerini saha raporu biçiminde dışa aktarır."),
        ],
        header="Üst Çubuk Göstergelerinin Anlamı",
    )
    add_heading(doc, "Saha Girdilerinin Oluşturulması", 2, "4.2")
    add_figure(
        doc,
        SCREENSHOTS / "02-senaryo-girdileri.png",
        "Şekil 4. Bir şehir merkezi senaryosunda depremzede ve enkaz girdilerinin işaretlenmesi",
    )
    add_body(
        doc,
        "Operatör “Depremzede Ekle” modunda haritaya hedef noktaları bırakır; “Enkaz Seç” modunda bina ayak izlerini çökmüş yapı olarak işaretler. “Kaldır” modu hatalı hedef veya enkaz işaretini geri alır. En az üç depremzede ve bir enkaz bulunmadan analiz başlatılamaz. Bu eşik, üç terminal bölgesi üreten kümeleme yaklaşımının asgari girdi koşuludur."
    )
    add_heading(doc, "Harita Sembol Dili", 2, "4.3")
    add_matrix_table(
        doc,
        ["Görsel öğe", "Mühendislik anlamı"],
        [
            ("Sarı/kırmızı küçük noktalar", "Yardım bekleyen hedef kullanıcı/depremsede konumları"),
            ("Kırmızı bina poligonu", "Kullanıcı tarafından enkaz kabul edilen bina"),
            ("Gri bina poligonu", "Ayakta bina; görüş hattı engeli olabilir"),
            ("Mavi terminal işaretçisi", "Uydu terminalinin önerilen saha konumu"),
            ("Numaralı renkli IRS işaretçisi", "Sinyal yansıtıcı birimin önerilen montaj noktası"),
            ("Yeşil hat", "Geometrik olarak açık kabul edilen haberleşme bacağı"),
            ("Kesikli/kırmızı hat", "Engelli veya kısmi görüş hattı"),
            ("Camgöbeği kısa çizgi", "IRS montaj cephesinin normal yönü"),
        ],
        [2600, 6760],
        font_size=8.8,
    )

    add_heading(doc, "Deterministik Yerleşim Algoritması", 1, "5.")
    add_figure(
        doc,
        WORKING / "islem-hatti.png",
        "Şekil 5. ARES-Reflect deterministik yerleşim işlem hattı",
        source="Kaynak: src/lib/algorithm.js ve src/lib/scoring.js üzerinden oluşturulmuştur.",
    )
    add_heading(doc, "Sahne ve Engel Modeli", 2, "5.1")
    add_body(
        doc,
        "Bina veri kümesindeki geçerli ayak izleri normalize edilir. Kullanıcının enkaz olarak işaretlediği bina kimlikleri, ayakta bina engel listesinden çıkarılır. Böylece blokaj kümesi “tüm bina ayak izleri – enkaz olarak işaretlenen binalar” biçiminde kurulur. Enkaz noktaları terminal/IRS için güvenlik ve mesafe hesabında dikkate alınır; ancak ayakta bir bina gibi sinyal engeli sayılmaz."
    )
    add_heading(doc, "Deterministik K-means Kümeleme", 2, "5.2")
    add_body(
        doc,
        "Depremzedeler üç hedef bölgesine ayrılır. Başlangıç merkezleri rastgele seçilmez: önce genel merkeze en uzak nokta, sonra mevcut merkezlere toplamda en uzak noktalar tercih edilir. Eşitlikler koordinat sırasıyla çözülür ve en fazla 50 iterasyon uygulanır. Bu nedenle aynı sahne aynı kümeleri üretir."
    )
    add_heading(doc, "Terminal Aday Uzayı", 2, "5.3")
    add_body(
        doc,
        "Her hedef kümesi için terminal adayları; küme çevresindeki 25 m aralıklı 7×7 ızgara, enkaz çevresindeki 18–90 m halkalar, enkaz–küme koridorları ve enkaz ile yakın ayakta bina arasındaki ara noktalar kullanılarak oluşturulur. Adayın yalnızca hedefe yakın olması yeterli değildir; işletilebilir bir bina–enkaz koridoru oluşturması ve geçerli IRS üretebilmesi gerekir."
    )
    add_kv_table(
        doc,
        [
            ("Bina/enkaz ayak izi", "Terminal hiçbir yapı ayak izinin içinde olamaz."),
            ("Enkaz kenar mesafesi", "8–75 m; 45 m’ye kadar olan bölge ideal kabul edilir."),
            ("Ayakta bina kenar mesafesi", "4–70 m; 35 m’ye kadar olan bölge ideal kabul edilir."),
            ("Karşılıklı yön açısı", "Enkaz ve ayakta bina yönleri arasında en az 125° aranır."),
            ("Terminal ayrımı", "100 m tercih edilir; 40 m altı katı çözüm kümesine alınmaz."),
        ],
        header="Terminal Saha Geçerliliği İçin Temel Sabitler",
    )
    add_heading(doc, "IRS Aday Üretimi ve Cepheye Oturtma", 2, "5.4")
    add_body(
        doc,
        "Her terminal için IRS adayları terminal–hedef doğrultusunda farklı oran ve ±45° yön sapmalarıyla üretilir. Temel arama uzaklığı 30–180 m’dir. Geçerli aday bulunamazsa daha geniş açılı ek tarama ve tek tek depremzede hedeflerine yönelik kapsamlı tarama uygulanır. Aday, 40 m içinde uygun bir bina cephesi varsa cephe dışına yaklaşık 1,5 m ötelenerek yerleştirilir; aksi durumda yaklaşık 8 m yüksekliğinde serbest direk modeli kullanılır."
    )
    add_heading(doc, "İki Bacaklı Fiziksel Eleme", 2, "5.5")
    add_body(
        doc,
        "IRS adayı iki ayrı bacak üzerinden değerlendirilir: Terminal → IRS ve IRS → hedef. Terminal–IRS hattını ayakta bina kesiyorsa aday doğrudan elenir. IRS–hedef tarafında ise her depremzede ayrı kontrol edilir; en az bir açık hedef varsa kısmi görüşlü fakat geçerli bir aday oluşabilir. Terminale aşırı yakın (kurulum ve yakın alan çakışması yaratacak kadar), sıfır açık kapsamalı veya kalite puanı %55’in altında kalan adaylar sonuç listesine alınmaz."
    )
    add_callout(
        doc,
        "Minimum yeterli IRS ilkesi",
        "Sistem terminal başına rastgele veya sabit üç IRS önermez. Önce erişilebilecek kapsama tavanını bulur; sonra aynı kapsama tavanına ulaşan en küçük IRS setini seçer. İkinci veya üçüncü IRS yalnızca yeni açık kapsama sağlıyorsa eklenir.",
        fill=LIGHT_CYAN,
        accent=CYAN,
    )

    add_heading(doc, "Puanlama ve Metriklerin Mühendislik Anlamı", 1, "6.")
    add_heading(doc, "IRS Mutlak Kalite Fonksiyonu", 2, "6.1")
    add_body(
        doc,
        "IRS kalite puanı aday kümesi içindeki göreli min–maks normalizasyonuna değil, [0,1] aralığına sınırlandırılmış fiziksel alt ölçütlere dayanır. Temel bileşik puan aşağıdaki biçimde uygulanır:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_run_text(
        p,
        "Q = 0,25 Lₜᵢ + 0,25 Lᵢₕ + 0,15 D + 0,10 H + 0,10 F + 0,10 G + 0,05 C − P",
        bold=True,
        color=NAVY,
        size=12,
    )
    add_matrix_table(
        doc,
        ["Terim", "Ağırlık", "Açıklama"],
        [
            ("Lₜᵢ", "%25", "Terminal–IRS görüş hattı oranı"),
            ("Lᵢₕ", "%25", "IRS–hedef açık görüş oranı"),
            ("D", "%15", "Toplam yol uzunluğunun uygunluğu"),
            ("H", "%10", "Montaj yüksekliği puanı"),
            ("F", "%10", "Cephe normalinin terminal/hedef/uydu yönlerine uyumu"),
            ("G", "%10", "Normalize edilmiş kestirimsel link kazancı"),
            ("C", "%5", "Açık kapsanan depremzede oranı"),
            ("P", "Ceza", "Blokaj başına 0,06 ve Fresnel ihlali için 0,20"),
        ],
        [1000, 1200, 7160],
        font_size=8.8,
    )
    add_body(
        doc,
        "Ek çarpanlar, enkaz alanı içinde/çok yakınında kalan veya yansıma verimi düşük olan adayların puanını düşürür. Bu fonksiyon bir RF sertifikasyon hesabı değil, geometrik açıdan uygulanabilir adayların aynı ölçekte karşılaştırılması için geliştirilmiş mühendislik sezgiselidir."
    )
    add_heading(doc, "Tahmini Link Kazancı", 2, "6.2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    add_run_text(
        p,
        "GdB = 12η + 10 log₁₀(L) − 1,5(d₁ + d₂)/100",
        bold=True,
        color=PURPLE,
        size=12,
    )
    add_body(
        doc,
        "Burada η yansıma verimini, L iki bacağın en düşük görüş oranını, d₁ + d₂ ise toplam yol uzunluğunu temsil eder. Formül; panel alanı, frekans, anten kazancı, serbest uzay kaybı, atmosferik kayıp ve alıcı hassasiyeti gibi tam link bütçesi bileşenlerini içermez. Bu nedenle KTR’de “yerleşim adaylarını kıyaslayan tahmini link kazancı” olarak adlandırılmalıdır."
    )
    add_heading(doc, "Terminal Puanı", 2, "6.3")
    add_matrix_table(
        doc,
        ["Bileşen", "Ağırlık", "İşlev"],
        [
            ("IRS set kalitesi", "%35", "Terminalin gerçekten destekleyebildiği IRS setinin kalitesi ve kapsaması"),
            ("Bina–enkaz saha uygunluğu", "%30", "Terminalin iki yapı arasındaki açık kurulum koridoru"),
            ("Açık görüş koridoru", "%15", "IRS koridoru ve doğrudan hedef görünürlüğü"),
            ("Yakınlık tabanlı kapsama", "%10", "Hedeflere uzaklık ve küme yoğunluğu"),
            ("Uydu erişim sezgiseli", "%5", "Enkaz doğrultularına göre sabit azimutlu erişim göstergesi"),
            ("Küme merkezine yakınlık", "%5", "Operasyonel hedef bölgesine mesafe"),
        ],
        [2700, 1100, 5560],
        font_size=8.7,
    )
    add_heading(doc, "Kalite Bantları", 2, "6.4")
    add_kv_table(
        doc,
        [
            ("%85–100 | Güçlü", "İki bacak ve kapsama birlikte güçlü; öncelikli saha adayı"),
            ("%70–84 | Uygun", "Geçerli ve kullanılabilir; saha marjı güçlü banda göre daha düşüktür"),
            ("%55–69 | Sınırda", "Geçerli fakat uzun yol, düşük yansıma veya sınırlı kapsama nedeniyle dikkat gerektirir"),
            ("%0–54 | Geçersiz", "Öneri kartında tutulmaz; fiziksel eşiklerden en az biri sağlanmamıştır"),
        ],
        header="Arayüzdeki Mutlak Kalite Yorumları",
    )

    add_heading(doc, "Çalışan Uygulamadan Örnek Senaryo", 1, "7.")
    add_body(
        doc,
        "Bu bölüm, ARES-Reflect’in canlı dağıtımında çalıştırılan gerçek bir senaryonun çıktısını inceler. Sahne, yoğun yapılaşmış bir kent merkezinde 6 depremzede ve 3 enkaz içerir. Yerel motor aynı sahne için saha puanları birbirine çok yakın üç terminal adayı (A: 90, B: 90, C: 89 / 100) üretmiş; her terminal için üç geçerli IRS adayı bulmuştur. Aşağıda örnek olarak Terminal A ayrıntılandırılmıştır. Bu değerler senaryoya özeldir ve genel sistem performansı olarak yorumlanmamalıdır."
    )
    add_heading(doc, "Haritada Terminal ve IRS Sonucunun Okunması", 2, "7.1")
    add_figure(
        doc,
        WORKING / "07a-sonuc-tam.png",
        "Şekil 6. Canlı uygulamada üretilen tam sonuç görünümü: solda harita yerleşimi, sağda analiz paneli",
        width=6.7,
        source="Kaynak: ARES-Reflect canlı dağıtımından (hareketli-uydu-teknolojileri.vercel.app) 19 Haziran 2026 tarihinde alınan ekran görüntüsü.",
    )
    add_body(
        doc,
        "Haritada mavi işaretçi terminalin önerilen konumudur ve merkezdeki enkaz kümesinin batısında yer alır. 1 (yeşil), 2 (turuncu) ve 3 (kırmızı) numaralı işaretçiler Terminal A için bulunan üç geçerli IRS adayını gösterir; üçü de merkezî enkazın çevresinde, terminal ile depremzede grupları arasındaki erişim koridorunda kümelenmiştir. Sarı noktalar yardım bekleyen depremzedeleri, kırmızı poligonlar enkaz olarak işaretlenen binaları temsil eder. Terminalden her IRS adayına ve adaylardan hedef gruplarına uzanan yeşil çizgiler, ayakta bina ile kesilmeyen iki bacaklı haberleşme koridorlarını gösterir; bir hattın yeşil olması o bacağın geometrik olarak açık olduğu anlamına gelir."
    )
    add_heading(doc, "Terminal Kartları ve IRS Karşılaştırması", 2, "7.2")
    add_figure(
        doc,
        WORKING / "07b-terminal-irs-paneli.png",
        "Şekil 7. Terminal sonuç kartları, Terminal A’nın üç geçerli IRS adayı ve seçim gerekçesi",
        width=5.4,
    )
    add_body(
        doc,
        "Karttaki 90/100 değeri terminalin bileşik saha puanıdır ve IRS kalite puanından farklıdır. “3 geçerli IRS” ifadesi, yerel motorun Terminal A için fiziksel eşiklerin tamamını sağlayan üç aday bulduğunu belirtir. Karşılaştırma tablosunda her aday; Terminal→IRS görüşü, IRS→hedef oranı, açık kapsanan depremzede sayısı, tahmini kazanç ve kalite puanı sütunlarıyla aynı ölçekte değerlendirilir."
    )
    add_matrix_table(
        doc,
        ["IRS adayı", "Terminal→IRS", "IRS→Hedef", "Açık kapsama", "Tahmini kazanç", "Kalite", "Durum"],
        [
            ("A-IRS-1", "%100", "%100", "6 / 6", "+5,6 dB", "%91", "Seçildi"),
            ("A-IRS-2", "%100", "%100", "6 / 6", "+6,6 dB", "%91", "Geçerli alternatif"),
            ("A-IRS-3", "%100", "%100", "6 / 6", "+4,4 dB", "%90", "Geçerli alternatif"),
        ],
        [1260, 1340, 1180, 1320, 1480, 900, 1880],
        font_size=8.2,
    )
    add_callout(
        doc,
        "Üç aday, tek seçim: minimum yeterli IRS ilkesinin somut görünümü",
        "Üç adayın da iki bacağı tam açıktır (%100 / %100) ve her biri tek başına altı depremzedenin tamamını kapsar. Kapsama tavanı tek bir IRS ile dolduğundan sistem üç birim önermez; aynı tavana ulaşan en küçük seti, yani tek IRS’yi seçer. A-IRS-1 ile A-IRS-2 kalite puanında eşit (%91), hatta A-IRS-2’nin tahmini kazancı daha yüksek olsa da seçim 53,1 m’lik en kısa toplam yola sahip A-IRS-1’den yana yapılmıştır. Kapsama ve kalite eşitlendiğinde daha kısa toplam yolun tercih edilmesi, daha düşük yol kaybı ve daha basit kurulum lehine deterministik bir eşitlik çözümüdür.",
        fill=LIGHT_CYAN,
        accent=CYAN,
    )
    add_heading(doc, "Seçilen IRS Biriminin Okunması", 2, "7.3")
    add_figure(
        doc,
        WORKING / "07c-irs-detay.png",
        "Şekil 8. Seçilen A-IRS-1 biriminin özet kartı ve operasyonel okuması",
        width=6.3,
    )
    add_kv_table(
        doc,
        [
            ("Terminal → IRS", "%100; ilk bacak ayakta bina ile kesilmez, tam açıktır."),
            ("IRS → Hedef", "%100; altı depremzedenin tamamına açık ikinci bacak kurulur."),
            ("Kalite puanı", "%91; “güçlü” kalite bandındadır (≈%85 ve üzeri)."),
            ("Terminale uzaklık", "17,3 m; terminal–IRS bacağının uzunluğudur."),
            ("Toplam sinyal yolu", "53,1 m; iki bacağın toplam kestirimsel uzunluğudur."),
            ("Tahmini link kazancı", "+5,6 dB; yerleşim adaylarını kıyaslayan karşılaştırma metriğidir."),
            ("Montaj", "Güney bina cephesi, yaklaşık 10 m yükseklik."),
        ],
        header="A-IRS-1 Sonucunun Yorumu",
    )
    add_body(
        doc,
        "Detay kartı, seçilen IRS biriminin açık hat durumunu, kapsamasını, terminale uzaklığını ve tahmini kazancını tek bakışta okunur biçimde sunar. A-IRS-1 için iki bacağın da tam açık (%100/%100) olması ve altı depremzedenin tamamının kapsanması, terminal–IRS–hedef koridorunun engelsiz kurulabildiğini gösterir. Kazanç değeri tam RF link bütçesi değil; görüş oranı, yansıma verimi ve toplam yol uzunluğundan türetilen kestirimsel bir karşılaştırma ölçütüdür."
    )

    add_heading(doc, "Proje Detayları Sayfası ve Bilgi Katmanı", 1, "8.")
    add_figure(
        doc,
        SCREENSHOTS / "06a-proje-detaylari-ust.png",
        "Şekil 9. Proje vizyonu ve temel tasarım hedefleri",
    )
    add_body(
        doc,
        "Proje Detayları sayfası hesaplama sonucu üretmez; uygulamanın teknik kavramlarını kullanıcıya açıklayan bilgi katmanıdır. 3GPP Release 17/18 doğrudan cihaza NTN vizyonu, 5–10 dB marjinal kazanç hedefi, 20 kg/140 W kısıtı, ±8° bozucu bastırma ve 8 saniyenin altında yeniden kilitlenme hedefleri bu sayfada ÖTR bağlamıyla sunulur."
    )
    add_figure(
        doc,
        SCREENSHOTS / "07-sistem-islem-akisi.png",
        "Şekil 10. Uygulama içinde görev tanımı ve saha karar akışının açıklanması",
    )
    add_body(
        doc,
        "Bu bilgi sayfasındaki açıklamalar, ekran metriklerinin yanlış yorumlanmasını önlemeyi amaçlar. Özellikle konumların bir dil modeli tarafından üretilmediği, bina blokajı ve geometrik kurallarla belirlendiği vurgulanır."
    )

    add_heading(doc, "Gemini Destek Katmanı", 1, "9.")
    add_heading(doc, "Yerel-Öncelikli Çalışma Sırası", 2, "9.1")
    add_bullet(doc, "Yerel motor kümeleri, terminal adaylarını, IRS adaylarını ve fiziksel geçerliliği hesaplar.")
    add_bullet(doc, "Sonuçlar Gemini yanıtı beklenmeden kullanıcıya gösterilir.")
    add_bullet(doc, "Gemini mevcut sayısal verilerden daha akıcı IRS açıklamaları üretir.")
    add_bullet(doc, "Gemini yalnızca yerel motorun geçerli bulduğu terminal kimliklerini yeniden sıralayabilir.")
    add_bullet(doc, "İsteğe bağlı görsel doğrulama harita görüntüsünü inceler; mühendislik kararının yerine geçmez.")
    add_heading(doc, "Yetki Sınırları", 2, "9.2")
    add_matrix_table(
        doc,
        ["Gemini yapabilir", "Gemini yapamaz"],
        [
            ("Sayısal sonucu doğal Türkçe ile açıklamak", "Yeni terminal veya IRS koordinatı üretmek"),
            ("Geçerli terminal kartlarının sunum sırasını değiştirmek", "Bloklu bir hattı geçerli hale getirmek"),
            ("Harita görüntüsünde ikincil tutarlılık denetimi yapmak", "Kalite eşiklerini veya yerel ağırlıkları değiştirmek"),
            ("Başarısız olduğunda yerel açıklamaya geri dönmek", "Yerel geometri sonucunu iptal etmek"),
        ],
        [4680, 4680],
        font_size=9,
    )
    add_callout(
        doc,
        "KTR’de kullanılacak doğru ifade",
        "“Yapay zekâ, fiziksel yerleşim kararını veren çekirdek algoritma değildir. Konumlar deterministik geometri motoruyla üretilmekte; üretken model yalnızca açıklama, eşdeğer geçerli sonuçların sunum sırası ve ikincil görsel denetim işlevlerinde kullanılmaktadır.”",
        fill=LIGHT_CYAN,
        accent=CYAN,
    )

    add_heading(doc, "Doğrulama ve Regresyon Testleri", 1, "10.")
    add_body(
        doc,
        "Yerleşim regresyon betiği 19 Haziran 2026 tarihinde altı deterministik senaryoda çalıştırılmıştır. Testler; üç terminal üretimi, terminalin bina içinde olmaması, bina–enkaz koridor mesafeleri, terminaller arası ayrım, IRS’nin terminalden asgari açıklığı, terminal–IRS ve IRS–hedef hatlarının ayakta bina kesmemesi, açık hedef bulunması ve farklı terminallerin IRS çakışmalarını denetlemektedir."
    )
    add_matrix_table(
        doc,
        ["Senaryo", "Terminal sayısı", "IRS dağılımı", "Test sonucu"],
        [
            ("Basit Kavşak", "3", "1 / 1 / 1", "Başarılı"),
            ("Çarşı Bölgesi", "3", "1 / 1 / 1", "Başarılı"),
            ("Okul Bölgesi", "3", "1 / 1 / 1", "Başarılı"),
            ("Şehir Merkezi", "3", "1 / 1 / 1", "Başarılı"),
            ("Hastane Bölgesi", "3", "1 / 1 / 1", "Başarılı"),
            ("Tam Kentsel Felaket", "3", "1 / 1 / 1", "Başarılı"),
        ],
        [2600, 1500, 2100, 3160],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Test sonucunun sınırı",
        "failureCount = 0 sonucu, kodda tanımlanmış geometrik kuralların bu senaryolarda ihlal edilmediğini kanıtlar. Fiziksel anten kazancı, gerçek saha radyo performansı veya mekanik dayanım için tek başına doğrulama değildir.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )

    add_heading(doc, "Varsayımlar, Sınırlamalar ve Gelecek Doğrulamalar", 1, "11.")
    limitations = [
        ("İki boyutlu LoS", "Blokaj hesabı bina ayak izi kesişimine dayanır; bina yüksekliğiyle 3B ışın izleme yapılmaz.", "3B şehir modeli ve yükseklik profili"),
        ("Fresnel modeli", "Mevcut sürümde ayrıntılı Fresnel zon yarıçapı hesaplanmaz.", "Frekans ve yol uzunluğuna bağlı 1. Fresnel zonu"),
        ("Link kazancı", "Kestirimsel aday karşılaştırma metriğidir.", "Frekans, EIRP, G/T, atmosferik kayıp ve alıcı hassasiyeti içeren bütçe"),
        ("Uydu yönü", "Sabit azimutlu sezgisel erişim puanı kullanılır.", "Uydu efemerisi, konum ve zamana bağlı azimut/elevasyon"),
        ("Bina yüksekliği", "Veri yoksa 12 m varsayılır.", "Lidar/DSM veya doğrulanmış OSM yükseklikleri"),
        ("Montaj güvenliği", "Cephe geometrisi önerilir; ankraj/taşıyıcılık hesaplanmaz.", "Saha keşfi, statik analiz ve iş güvenliği prosedürü"),
        ("RF yansıma", "Simetrik açı ve basitleştirilmiş verim modeli kullanılır.", "Panel hücre faz modeli ve ölçüm tabanlı yansıma diyagramı"),
        ("Gemini doğrulaması", "Görsel ve tavsiye niteliğindedir.", "Mühendislik kabul testi olarak kullanılmamalıdır"),
    ]
    add_matrix_table(
        doc,
        ["Konu", "Mevcut yaklaşım", "KTR/prototip aşamasında önerilen doğrulama"],
        limitations,
        [1900, 3560, 3900],
        font_size=8.2,
    )
    add_heading(doc, "ÖTR’de KTR Öncesi Netleştirilmesi Gereken Noktalar", 2, "11.1")
    add_bullet(doc, "Ana kontrolcü metinde ESP32, tedarik takviminde STM32 olarak geçmektedir; nihai kontrolcü seçimi tekilleştirilmelidir.")
    add_bullet(doc, "Takvimde “adaptif PID”, teknik mimaride dış PI ve iç PID belirtilmektedir; kontrol algoritması ve katsayıları netleştirilmelidir.")
    add_bullet(doc, "Toplam kütle, alt sistem güç tüketimi ve emniyet payları ölçülmüş bütçe tablosuyla verilmelidir.")
    add_bullet(doc, "EKF durum/ölçüm modeli, örnekleme frekansı, gürültü kovaryansları ve deney sonuçları raporlanmalıdır.")
    add_bullet(doc, "5–10 dB hedefi, simülasyon veya ölçüm koşullarıyla birlikte gösterilmeli; yazılımın sezgisel link_gain_db metriğiyle karıştırılmamalıdır.")

    add_heading(doc, "KTR Yazımında Kullanılabilecek Teknik İfadeler", 1, "12.")
    expressions = [
        ("Sistem tanımı", "ARES-Reflect, afet sahasında terminal ve IRS yerleşimini bina ayak izi tabanlı görüş hattı, kapsama ve kurulum koridoru ölçütleriyle belirleyen deterministik bir karar destek sistemidir."),
        ("Deterministik yapı", "Başlangıç merkezi seçiminde rassal örnekleme kullanılmadığından aynı saha girdileri aynı yerel yerleşim sonucunu üretmektedir."),
        ("İki bacaklı kanal", "Her IRS adayı terminal–IRS ve IRS–hedef bacaklarında ayrı ayrı blokaj kontrolüne tabi tutulmakta; ilk bacağı ayakta bina ile kesilen adaylar doğrudan elenmektedir."),
        ("IRS adedi", "IRS sayısı sabit bir kota değildir; ulaşılabilir kapsama tavanını sağlayan en küçük geçerli set seçilmektedir."),
        ("Yapay zekâ rolü", "Üretken yapay zekâ koordinat veya fiziksel geçerlilik üretmemekte, yalnızca yerel hesapların açıklanması ve sunum sırasının desteklenmesi amacıyla kullanılmaktadır."),
        ("Kazanç metriği", "Uygulamadaki link kazancı, tam RF link bütçesi değil; görüş oranı, yansıma verimi ve toplam yol uzunluğundan türetilen kestirimsel bir karşılaştırma ölçütüdür."),
        ("Doğrulama", "Regresyon testleri tanımlı geometrik kısıtların ihlal edilmediğini doğrulamakta; fiziksel prototip performansı için ayrıca RF ve mekanik deney gerekmektedir."),
    ]
    add_kv_table(doc, expressions, header="Doğrudan Kullanılabilir Akademik İfade Örnekleri")

    add_heading(doc, "Veri Sözlüğü", 1, "13.")
    glossary = [
        ("quality_score", "IRS’nin 0–1 aralığındaki mutlak bileşik kalite puanı; arayüzde yüzdeye çevrilir."),
        ("term_los", "Terminal–IRS bacağının açıklık oranı."),
        ("vic_los", "IRS’den hedef kümesine açık erişim oranı."),
        ("survivors_covered_clear", "Ayakta bina ile kesilmeyen hat üzerinden erişilen depremzede sayısı."),
        ("total_path_m", "Terminal–IRS ve IRS–hedef yollarının toplam uzunluğu."),
        ("link_gain_db", "Basitleştirilmiş yansıma/LoS/yol modelinden üretilen tahmini karşılaştırma kazancı."),
        ("reflection_efficiency", "Gelme–çıkış yönleri ile cephe geometrisinden türetilen yansıma uygunluğu."),
        ("facade_alignment", "Cephe normalinin terminal, hedef ve sabit uydu yönüyle bileşik uyumu."),
        ("validity_status", "Adayın valid, borderline veya invalid sınıfı."),
        ("constrained_reason", "Sınırda veya geçersiz adayın temel mühendislik nedeni."),
        ("selection_reason", "Seçilen IRS’nin kapsama/kalite/set katkısını açıklayan yerel gerekçe."),
        ("siteEvaluation", "Terminalin enkaz ve ayakta bina arasındaki kurulum koridoru değerlendirmesi."),
    ]
    add_kv_table(doc, glossary, widths=(2700, 6660), header="Uygulama Alanlarının Teknik Karşılığı")

    add_heading(doc, "Kaynaklar ve İzlenebilirlik", 1, "14.")
    refs = [
        "ARES-Reflect güncel kaynak kodu: src/lib/algorithm.js, geometry.js, scoring.js, buildings.js, gemini.js; src/hooks/useAnalysis.js ve arayüz bileşenleri.",
        "ARES-Reflect README.md, yerel-öncelikli işlem hattı ve kullanıcı iş akışı açıklamaları.",
        "Hareketli Uydu Terminali PATH ÖTR dokümanı; fiziksel terminal, RIS/IRS, EKF, kaskad kontrol ve tasarım kısıtları.",
        "scripts/placement-regression.mjs; altı deterministik senaryoya ait geometrik regresyon testleri.",
        "Bu belgedeki ekran görüntüleri çalışan ARES-Reflect uygulamasından (yerel geliştirme sunucusu ve canlı dağıtım) 19 Haziran 2026 tarihinde alınmıştır.",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    add_callout(
        doc,
        "Son kullanım notu",
        "KTR’ye aktarım yapılırken bu belgedeki ekran görüntüleri şekil olarak, akademik ifade örnekleri açıklama metni olarak kullanılabilir. Ancak yarışma şablonundaki bölüm sırası, sayfa limiti ve zorunlu format ayrıca uygulanmalıdır.",
        fill=LIGHT_CYAN,
        accent=CYAN,
    )


def finalize(doc):
    core = doc.core_properties
    core.title = "ARES-Reflect Teknik Sistem Açıklama Dosyası"
    core.subject = "KTR hazırlığı için teknik kaynak belge"
    core.author = "ARES-Reflect Proje Ekibi"
    core.keywords = "ARES-Reflect, IRS, RIS, uydu terminali, afet haberleşmesi, KTR"
    core.comments = "Güncel kaynak kod, çalışan uygulama, regresyon testleri ve ÖTR temel alınmıştır."

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.88)
        section.right_margin = Inches(0.88)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    make_assets()
    document = setup_document()
    add_cover(document)
    add_contents(document)
    add_main_content(document)
    finalize(document)
