from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SHOTS = ROOT / "artifacts" / "yeni-rapor" / "screenshots"
STATUS_SHOT = Path(
    r"C:\Users\oguzh\AppData\Local\Temp\codex-clipboard-0bf7c342-9e07-48be-b676-829396cb23ec.png"
)
SCENARIO_SHOT = Path(
    r"C:\Users\oguzh\AppData\Local\Temp\codex-clipboard-bb6cc022-eb9f-452c-9379-aac0a171ec79.png"
)
OUT = ROOT / "output" / "docx" / "ARES-Reflect_Akademik_Arayuz_ve_Teknik_Mimari.docx"

NAVY = "142033"
CYAN = "10B8C8"
BLUE = "2F6FFF"
PURPLE = "8B5CF6"
GREEN = "16A36A"
AMBER = "C68A00"
INK = "1F2937"
MUTED = "667085"
PALE = "F3F6FA"
PALE_CYAN = "EAFBFD"
WHITE = "FFFFFF"


def rgb(value):
    return RGBColor.from_string(value)


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=140, start=180, bottom=140, end=180):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width_dxa))
    grid.append(col)
    for row in table.rows:
        cell = row.cells[0]
        tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tc_w)
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")
        set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_font(run, size=8.5, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ARES-Reflect  |  Arayüz ve Sistem Çalışma Akışı")
    set_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)
    return doc


def add_body(doc, text, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r, size=10.7)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}[level], color={1: NAVY, 2: BLUE, 3: NAVY}[level], bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=10.4)
    return p


def add_callout(doc, title, text, fill=PALE_CYAN, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=10.3, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    r2 = p2.add_run(text)
    set_font(r2, size=9.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_screenshot_path(doc, image_path, caption, width=6.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(8)
    cr = cp.add_run(caption)
    set_font(cr, size=8.8, color=MUTED, italic=True)


def add_screenshot(doc, filename, caption, width=6.55):
    add_screenshot_path(doc, SHOTS / filename, caption, width)


def new_page(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section = doc.sections[-1]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(42)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("TEKNOFEST UYDU HABERLEŞME PROJESİ")
    set_font(kr, size=10.5, color=CYAN, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run("ARES-Reflect")
    set_font(tr, size=31, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    sr = subtitle.add_run("Arayüz ve Sistem Çalışma Akışı")
    set_font(sr, size=18, color=BLUE, bold=True)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.45)
    lead.paragraph_format.right_indent = Inches(0.45)
    lead.paragraph_format.space_after = Pt(28)
    lr = lead.add_run(
        "Bu dosya, uygulamanın light mod arayüzlerini ve terminal–IRS yerleşim kararının "
        "ekranda nasıl üretildiğini; LoS/NLoS, iki bacaklı röle kanalı, blokaj ve kestirimsel "
        "link bütçesi kavramlarıyla kısa, teknik ve görsel bir akışta açıklar."
    )
    set_font(lr, size=12, color=MUTED)

    add_callout(
        doc,
        "ÇEKİRDEK MÜHENDİSLİK PRENSİBİ",
        "Terminal ve IRS koordinatları yerel deterministik geometri motoru tarafından hesaplanır. "
        "Her Terminal→IRS ve IRS→hedef bağlantısı bağımsız LoS/NLoS denetimine alınır. Üretken "
        "yapay zekâ fiziksel geçerlilik üretmez; yalnızca açıklama, geçerli sonuçların sunum sırası "
        "ve ikincil görsel kontrol görevlerinde kullanılır.",
        fill=PALE_CYAN,
        accent=CYAN,
    )

    add_heading(doc, "Sistem akışı", 2)
    steps = [
        "Saha girdileri: depremzede noktaları, enkaz binaları ve sağlam bina ayak izleri alınır.",
        "Depremzedeler deterministik olarak üç operasyon bölgesine ayrılır.",
        "Terminal adayları enkaz güvenliği, bina mesafesi ve açık koridor ölçütleriyle taranır.",
        "IRS adayları sağlam bina cephelerinde üretilir; iki bağlantı ayağı LoS, PARTIAL_NLoS ve FULL_NLoS durumlarıyla denetlenir.",
        "Geçerli terminal–IRS setleri kapsama, yol uzunluğu, cephe uyumu, kestirimsel link marjı ve kalite puanıyla sıralanır.",
        "Sonuç harita, terminal kartları, IRS tablosu ve teknik ayrıntı penceresinde gösterilir.",
    ]
    for step in steps:
        add_bullet(doc, step)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Haziran 2026")
    set_font(r, size=9.5, color=MUTED, bold=True)


def add_main_pages(doc):
    new_page(doc)
    add_heading(doc, "1. İşlem kaynağı ve sistem durum göstergeleri", 1)
    add_screenshot_path(
        doc,
        STATUS_SHOT,
        "Şekil 1. Yerleşim, açıklama, yeniden sıralama ve görsel doğrulama katmanlarının anlık durum göstergeleri.",
        width=6.55,
    )
    add_body(
        doc,
        "Üst durum çubuğu, karar zincirindeki her fonksiyonun hangi hesaplama katmanından beslendiğini "
        "operatöre bildirir. Bu ayrım özellikle emniyet-kritik yerleşim kararının üretken modelden bağımsız "
        "olduğunu ve fiziksel LoS/NLoS doğrulamasının yerel geometri motorunda tutulduğunu göstermek için kullanılır."
    )
    add_bullet(
        doc,
        "Yerleşim: Yerel — terminal ve IRS koordinatları, bina ayak izleri, montaj yükseklikleri ve iki bacaklı kanal blokajı kullanılarak deterministik biçimde hesaplanır.",
    )
    add_bullet(
        doc,
        "Açıklama: Yerel — IRS seçim gerekçesi; açık kapsama, kalite puanı, toplam yol ve NLoS durumu gibi sayısal çıktılardan şablon tabanlı olarak üretilir.",
    )
    add_bullet(
        doc,
        "Rerank: Yerel — terminal kartlarının sıralaması yerel bileşik saha puanına göre korunur; Gemini etkinleşse bile yalnızca fiziksel olarak geçerli yerel çözümler yeniden sıralanabilir.",
    )
    add_bullet(
        doc,
        "Doğrulama: Hazır — yerleşim tamamlandıktan sonra harita görüntüsünün ikincil görsel denetime gönderilebileceğini belirtir; bu durum RF veya LoS kabul testi değildir.",
    )
    add_callout(
        doc,
        "Arayüz kontrolü",
        "Dark düğmesi yalnızca görsel temayı değiştirir. PDF Dışa Aktar işlevi ise terminal, IRS, LoS/NLoS "
        "oranları, kapsama ve kestirimsel link kazancı gibi mevcut analiz verilerini rapor çıktısına dönüştürür; "
        "hesaplama sonucunu değiştirmez.",
        fill=PALE,
        accent=BLUE,
    )

    new_page(doc)
    add_heading(doc, "2. Deterministik demo senaryoları", 1)
    add_screenshot_path(
        doc,
        SCENARIO_SHOT,
        "Şekil 2. Enkaz ve depremzede yoğunluğu kademeli olarak artan altı deterministik demo senaryosu.",
        width=6.35,
    )
    add_body(
        doc,
        "Demo Senaryolar menüsü, operatörün veri girişini hızlandıran ve algoritmanın farklı yoğunluk ile "
        "blokaj koşullarındaki davranışını karşılaştırmalı olarak incelemeye yarayan ön tanımlı saha "
        "konfigürasyonlarıdır. Senaryolar basit, orta, zor ve aşırı sınıfları altında artan enkaz sayısı, "
        "depremzede yoğunluğu ve muhtemel NLoS koridor karmaşıklığı içerir."
    )
    add_bullet(doc, "Senaryo 1–2: düşük düğüm yoğunluğu ve sınırlı bina blokajı altında temel terminal–IRS yerleşimini sınar.")
    add_bullet(doc, "Senaryo 3: orta yoğunlukta kapsama, terminal ayrışması ve cephe seçimi dengesini görünür kılar.")
    add_bullet(doc, "Senaryo 4–5: artan enkaz ve kapalı yol koşullarında LoS kaybı, PARTIAL_NLoS ve alternatif IRS cephelerini zorlar.")
    add_bullet(doc, "Senaryo 6: 24 depremzede ve 6 enkaz ile aday uzayını, üç terminal/üç IRS koşulunu ve yoğun kentsel blokajı stres altında değerlendirir.")
    add_callout(
        doc,
        "Doğrulama sınırı",
        "Demo sınıfları yazılımın deterministik geometri ve seçim akışını tekrarlanabilir biçimde göstermek "
        "içindir. Bunlar tek başına saha RF ölçümü, 3B ışın izleme, Fresnel zonu veya fiziksel prototip kabul testi değildir.",
        fill="FFF8E8",
        accent=AMBER,
    )

    new_page(doc)
    add_heading(doc, "3. Saha girdilerinin oluşturulması", 1)
    add_screenshot(
        doc,
        "02-saha-girdileri.png",
        "Şekil 3. Örnek senaryoda 12 depremzede ve 4 enkaz binasının harita üzerinde işaretlenmesi.",
    )
    add_body(
        doc,
        "Sarı işaretçiler yardım bekleyen cihaz veya kullanıcı konumlarını, kırmızı bina alanları "
        "enkaz olarak kabul edilen yapıları gösterir. Haritadaki diğer bina ayak izleri sağlam yapı "
        "olarak değerlendirilir ve iki bacaklı röle kanalının LoS/NLoS blokaj kontrolünde kullanılır."
    )
    add_bullet(doc, "Enkaz olarak seçilen yapılar, haberleşme hattını kesen sağlam bina listesine alınmaz.")
    add_bullet(doc, "Sağlam bina geometrileri terminal montajı, IRS cephe üretimi ve yükseklik duyarlı LoS/NLoS denetiminde kullanılır.")
    add_bullet(doc, "Aynı saha girdileri, rassal başlangıç kullanılmadığı için aynı yerel yerleşim sonucunu üretir.")

    new_page(doc)
    add_heading(doc, "4. Terminal ve IRS yerleşim sonucunun okunması", 1)
    add_screenshot(
        doc,
        "03-terminal-ve-irs-sonuclari.png",
        "Şekil 4. Analiz sonrası terminal konumları, IRS bağlantı çizgileri ve karşılaştırma tablosu.",
    )
    add_body(
        doc,
        "Yerel motor depremzede noktalarını üç kümeye ayırır ve her küme için terminal adaylarını "
        "tarar. Terminal konumu yalnızca kümeye yakınlıkla seçilmez; enkazdan güvenli uzaklık, sağlam "
        "bina ilişkisi, açık koridor ve destekleyebildiği IRS setinin kalitesi birlikte değerlendirilir."
    )
    add_bullet(doc, "Her terminal için üç farklı sağlam bina cephesinde üç geçerli IRS bulunması gerekir.")
    add_bullet(doc, "Terminal→IRS ve IRS→hedef hatları ayrı ayrı yükseklik duyarlı LoS/NLoS blokaj kontrolünden geçer.")
    add_bullet(doc, "FULL_NLoS durumunda açık hedef kalmayan aday elenir; PARTIAL_NLoS durumunda yalnızca açık kalan hedefler kapsama hesabına alınır.")
    add_bullet(doc, "Tablodaki açık kapsama, kestirimsel link kazancı ve kalite değerleri aynı terminale ait IRS adaylarını karşılaştırır.")
    add_callout(
        doc,
        "Harita çizgileri",
        "Haritadaki bağlantılar yalnızca görsel süs değildir. Çizgiler, optimizasyon sırasında kullanılan "
        "bina blokaj kurallarıyla yeniden kontrol edilen terminal–IRS–hedef LoS/NLoS koridorlarını temsil eder.",
        fill=PALE,
        accent=BLUE,
    )

    new_page(doc)
    add_heading(doc, "5. IRS ayrıntı penceresi", 1)
    add_screenshot(
        doc,
        "04-irs-ozet-penceresi.png",
        "Şekil 5. Seçilen IRS biriminin kalite, mesafe, kapsama ve montaj bilgilerinin özeti.",
        width=6.35,
    )
    add_body(
        doc,
        "IRS ayrıntı penceresi, seçilen yansıtıcı yüzeyin sahada neden kullanılabileceğini tek ekranda "
        "özetler. Terminal→IRS ve IRS→hedef oranları iki bağlantı ayağının LoS oranını; terminale uzaklık "
        "ve toplam sinyal yolu ise IRS üzerinden kurulan iki-atlamalı kırık hattın geometrik uzunluğunu gösterir."
    )
    add_bullet(doc, "Hat durumu CLEAR, PARTIAL_NLoS veya FULL_NLoS olarak sınıflandırılır; kalite puanı bu fiziksel durumdan bağımsız yorumlanmaz.")
    add_bullet(doc, "Kalite puanı; LoS oranı, toplam yol, montaj yüksekliği, cephe hizası, tahmini link kazancı ve kapsama bileşenlerinden oluşur.")
    add_bullet(doc, "Montaj alanı bina cephesi ve yaklaşık yükseklik bilgisiyle birlikte verilir.")
    add_bullet(doc, "Tahmini link kazancı tam bir RF link bütçesi değil, adayları aynı model içinde kıyaslayan kestirimsel metriktir.")
    add_callout(
        doc,
        "Fiziksel eleme",
        "Bir IRS adayı, terminalden panele veya panelden hedeflere açık bir LoS bağlantısı kuramıyorsa "
        "yüksek sunum puanıyla geçerli hale getirilemez; fiziksel kısıtlar önceliklidir.",
        fill="FFF8E8",
        accent=AMBER,
    )

    new_page(doc)
    add_heading(doc, "6. İkincil görsel doğrulama", 1)
    add_screenshot(
        doc,
        "06-gorsel-dogrulama.png",
        "Şekil 6. Hesaplanan yerleşimin harita görüntüsü üzerinden yapılan ikincil görsel kontrolü.",
        width=6.35,
    )
    add_body(
        doc,
        "Yerleşim tamamlandıktan sonra harita görüntüsü isteğe bağlı olarak görsel doğrulama katmanına "
        "gönderilebilir. Bu katman, terminal ve IRS işaretlerinin ekranda tutarlı görünüp görünmediğine "
        "ilişkin bir güven değerlendirmesi sunar."
    )
    add_callout(
        doc,
        "Yetki sınırı",
        "Görsel doğrulama yeni koordinat üretmez, NLoS bir bağlantıyı LoS kabul edemez ve yerel geometri motorunun "
        "kararını değiştiremez. Nihai fiziksel geçerlilik terminal–IRS ve IRS–hedef hat hesaplarına aittir.",
        fill=PALE_CYAN,
        accent=GREEN,
    )
    add_heading(doc, "Yerel motor ile yapay zekâ katmanının görev ayrımı", 2)
    add_bullet(doc, "Yerel motor: kümeleme, koordinat üretimi, cephe seçimi, blokaj kontrolü, puanlama ve nihai geçerlilik.")
    add_bullet(doc, "Yapay zekâ katmanı: teknik açıklamayı akıcılaştırma, yalnızca geçerli terminalleri yeniden sıralama ve görsel kontrol.")

    new_page(doc)
    add_heading(doc, "7. Proje bağlamı ve teknik özet ekranı", 1)
    add_screenshot(
        doc,
        "07-proje-detaylari.png",
        "Şekil 7. Uydu haberleşme hedefleri ile yazılım yerleşim motorunu aynı çerçevede sunan proje detay ekranı.",
    )
    add_body(
        doc,
        "Proje Detayları ekranı, arayüzde görülen yerleşim kararını hareketli uydu terminali ve IRS/RIS "
        "yaklaşımıyla ilişkilendirir. 3GPP Release 17/18 NTN vizyonu, 5–10 dB marjinal bağlantı kazancı hedefi, "
        "mekanik stabilizasyon ve yeniden kilitlenme gereksinimleri bu bağlamda sunulur. Bu sayfa hesaplama "
        "yapmaz; sistem hedeflerini ve yazılım motorunun NLoS erişim problemindeki rolünü teknik olarak özetler."
    )
    add_heading(doc, "Sonuç", 2)
    add_body(
        doc,
        "ARES-Reflect; afet sahasındaki kullanıcı, enkaz ve bina verilerini alarak terminal ile akıllı "
        "yansıtıcı yüzeylerin kurulabileceği konumları deterministik biçimde belirleyen bir karar destek "
        "arayüzüdür. Sistem, doğrudan terminal–hedef hattının NLoS kaldığı bölgelerde IRS/RIS destekli iki "
        "bacaklı bir röle koridoru arar; her öneriyi LoS oranı, blokaj, kapsama, cephe uyumu ve kestirimsel "
        "link bütçesi üzerinden denetleyerek sonucu açıklanabilir hale getirir."
    )
    add_callout(
        doc,
        "Raporlama için kullanılacak kısa teknik ifade",
        "“Doğrudan uydu/terminal erişiminin NLoS kaldığı afet bölgelerinde, IRS destekli iki-atlamalı kanal "
        "oluşturulmaktadır. Yerleşim koordinatları deterministik geometri motoru tarafından üretilmekte; "
        "Terminal→IRS ve IRS→hedef bağlantıları bina geometrileri ile montaj yükseklikleri kullanılarak ayrı "
        "ayrı LoS/NLoS denetimine alınmaktadır. Üretken yapay zekâ fiziksel kararın yerine geçmeden yalnızca "
        "açıklama, geçerli çözümlerin sunum sırası ve ikincil görsel kontrol katmanında görev almaktadır.”",
        fill=PALE_CYAN,
        accent=CYAN,
    )


def add_academic_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(38)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    set_font(kicker.add_run("TEKNOFEST HAREKETLİ UYDU TERMİNALİ PROJESİ"), size=10.5, color=CYAN, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("ARES-Reflect"), size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(
        subtitle.add_run("Arayüz Tabanlı Teknik Mimari ve Karar Destek İş Akışı"),
        size=17.5,
        color=BLUE,
        bold=True,
    )

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.4)
    lead.paragraph_format.right_indent = Inches(0.4)
    lead.paragraph_format.space_after = Pt(22)
    set_font(
        lead.add_run(
            "Afet sonrası NLoS haberleşme bölgelerinde akıllı yansıtıcı yüzey destekli erişim "
            "koridorlarının; coğrafi veri işleme, deterministik uzamsal optimizasyon ve "
            "yükseklik duyarlı görüş hattı denetimiyle oluşturulması"
        ),
        size=11.7,
        color=MUTED,
    )

    add_callout(
        doc,
        "SİSTEM SINIFI VE KAPSAM",
        "ARES-Reflect; OpenStreetMap bina geometrileri, afet düğümleri ve enkaz sınıflandırmasını "
        "aynı jeo-uzamsal sahne modelinde birleştiren, kural tabanlı ve deterministik bir karar "
        "destek yazılımıdır. Fiziksel koordinat üretimi yerel çözücüye aittir; üretken yapay zekâ "
        "karar otoritesi değil, açıklanabilirlik ve ikincil görsel inceleme katmanıdır.",
    )
    add_heading(doc, "Teknik işlem zinciri", 2)
    for text in (
        "Jeo-uzamsal veri normalizasyonu, bina poligon topolojisi ve enkaz/sağlam yapı ayrıştırması",
        "Deterministik K-Means ile hedef düğümlerin üç servis bölgesine bölünmesi",
        "Izgara, halka ve koridor örneklemesiyle terminal aday uzayının oluşturulması",
        "Cephe normal vektörü, azimut uyumu ve montaj yüksekliğiyle yansıtıcı yüzey aday üretimi",
        "Besleme ve yansıtılmış erişim bacaklarında CLEAR, PARTIAL_NLoS ve FULL_NLoS sınıflandırması",
        "Katı fiziksel kısıtlar ile çok ölçütlü kalite fonksiyonunun birlikte uygulanması",
    ):
        add_bullet(doc, text)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    set_font(p.add_run("Haziran 2026"), size=9.5, color=MUTED, bold=True)


def add_academic_pages(doc):
    new_page(doc)
    add_heading(doc, "1. Hesaplama kökeni, veri soyu ve durum telemetrisi", 1)
    add_screenshot_path(
        doc,
        STATUS_SHOT,
        "Şekil 1. Yerleşim çözücüsü, açıklama üreticisi, yeniden sıralama ve görsel doğrulama servislerinin durum telemetrisi.",
        width=6.55,
    )
    add_body(
        doc,
        "Üst durum şeridi, hesaplama zincirindeki işlevlerin veri kökenini (provenance) ve çalışma "
        "durumunu görünür kılar. “Yerleşim: Yerel” etiketi, koordinatların harici bir dil modeli "
        "tarafından üretilmediğini; deterministik kümeleme, aday tarama, poligon-kesişim testi ve "
        "kısıtlı optimizasyon sonucunda elde edildiğini bildirir. Bu ayrım, tekrar üretilebilirlik, "
        "denetlenebilirlik ve emniyet-kritik kararların izlenebilirliği açısından temel tasarım ilkesidir."
    )
    add_bullet(doc, "Açıklama servisi, sayısal özellik vektörünü doğal dil gerekçesine dönüştürür; geometri durumunu değiştirme yetkisi yoktur.")
    add_bullet(doc, "Yeniden sıralama servisi yalnızca uygunluk filtresini geçmiş çözüm kümesinde çalışır; katı kısıtları ihlal eden adaylar sıralama uzayına alınmaz.")
    add_bullet(doc, "Görsel doğrulama, render edilmiş haritanın semantik tutarlılığını inceler; RF kanal kestirimi veya elektromanyetik uygunluk analizi değildir.")
    add_bullet(doc, "PDF dışa aktarma, analiz anındaki durum vektörünü ve mühendislik metriklerini değişmez bir saha kayıt paketine dönüştürür.")
    add_callout(
        doc,
        "Yerel-öncelikli orkestrasyon",
        "Sonuç, çevrim içi yapay zekâ yanıtı beklenmeden oluşturulur. Bu mimari; ağ erişiminin "
        "kesintili olduğu afet sahalarında düşük gecikmeli ve çevrim dışı çalışmaya uygun bir "
        "fail-safe davranış sağlar.",
        fill=PALE,
        accent=BLUE,
    )

    new_page(doc)
    add_heading(doc, "2. Senaryo uzayı ve parametrik stres seviyeleri", 1)
    add_screenshot_path(
        doc,
        SCENARIO_SHOT,
        "Şekil 2. Düğüm yoğunluğu, enkaz sayısı ve kentsel blokaj karmaşıklığı artırılan deterministik deney senaryoları.",
        width=6.25,
    )
    add_body(
        doc,
        "Demo menüsü, algoritmik davranışın farklı saha yüklerinde karşılaştırılması için tanımlanmış "
        "parametrik deney kümelerini sunar. Zorluk seviyesi yalnızca depremzede sayısının artmasıyla "
        "değil; engel yoğunluğu, aday cephe sayısı, mekânsal yayılım, rota kapanması ve NLoS oluşma "
        "olasılığının birlikte yükselmesiyle belirlenir."
    )
    add_bullet(doc, "Basit sınıf: düşük düğüm yoğunluğu, seyrek engel alanı ve geniş uygulanabilir çözüm bölgesi.")
    add_bullet(doc, "Orta sınıf: artan küme varyansı, cephe seçimi belirsizliği ve kapsama çakışması.")
    add_bullet(doc, "Zor sınıf: daralan görünürlük koridorları, kısıtlı montaj yüzeyi ve daha yüksek yol sapması.")
    add_bullet(doc, "Aşırı sınıf: büyük aday uzayı, yoğun poligon-kesişim yükü ve üçlü IRS seti için kombinatoryal seçim baskısı.")
    add_callout(
        doc,
        "Deneysel geçerlilik sınırı",
        "Bu senaryolar yazılımın geometrik regresyonunu ve deterministikliğini değerlendirir. "
        "Fresnel bölgesi açıklığı, çok yollu sönümleme, gölgelenme dağılımı, atmosferik kayıp ve "
        "ölçüm tabanlı kanal kestirimi mevcut modelin dışında tutulmuştur.",
        fill="FFF8E8",
        accent=AMBER,
    )

    new_page(doc)
    add_heading(doc, "3. Jeo-uzamsal sahne modeli ve giriş verisi", 1)
    add_screenshot(
        doc,
        "02-saha-girdileri.png",
        "Şekil 3. Afet düğümleri, enkaz sınıfları ve OSM bina ayak izlerinden oluşan jeo-uzamsal sahne modeli.",
    )
    add_body(
        doc,
        "Harita katmanı; enlem-boylam koordinatlı afet düğümleri, kullanıcı tarafından enkaz olarak "
        "etiketlenen yapılar ve OpenStreetMap üzerinden alınan bina poligonlarını ortak bir sahne "
        "modelinde birleştirir. Kapalı poligonların merkez noktası, sınırlayıcı kutusu (bounding box), "
        "yaklaşık yarıçapı ve varsa bina yüksekliği/kat bilgisi ön işleme aşamasında normalize edilir."
    )
    add_bullet(doc, "Enkaz sınıfına alınan poligonlar yapısal engel listesinden çıkarılır; ayakta kalan yapılar obstrüksiyon geometrisi olarak korunur.")
    add_bullet(
        doc,
        "Geometrik uygunluk iki ayrı testle denetlenir: nokta-poligon testi, önerilen terminal veya "
        "IRS konumunun bina ayak izi içinde kalıp kalmadığını; doğru parçası-poligon testi ise "
        "besleme ve yansıtılmış erişim güzergâhlarının sağlam bina ayak izleriyle kesişip "
        "kesişmediğini belirler.",
    )
    add_bullet(doc, "Yükseklik verisi bulunmadığında kontrollü bir varsayılan değer kullanılır; bu nedenle sonuçlar ayrıntılı 3B şehir modeli yerine indirgenmiş 2.5B yaklaşımı temsil eder.")
    add_bullet(doc, "Deterministik başlangıç merkezleri ve bağlayıcı sıralama kuralları, aynı giriş sahnesinde aynı kümeleme sonucunun elde edilmesini sağlar.")

    new_page(doc)
    add_heading(doc, "4. Kısıtlı uzamsal optimizasyon ve aday seçimi", 1)
    add_screenshot(
        doc,
        "03-terminal-ve-irs-sonuclari.png",
        "Şekil 4. Çok ölçütlü amaç fonksiyonuyla seçilen terminal düğümleri, yansıtıcı yüzeyler ve röle koridorları.",
    )
    add_body(
        doc,
        "Hedef düğümler deterministik K-Means ile üç servis kümesine ayrılır. Her küme için centroid, "
        "enkaz çeperi, koridor doğrultusu ve yerel bina halkası çevresinde aday örnekleme yapılır. "
        "Çözücü; saha uygunluğu, doğrudan görünürlük potansiyeli, erişilebilirlik, kapsama oranı, "
        "yansıtıcı set kalitesi ve terminaller arası mekânsal ayrışmayı tek bir bileşik amaç fonksiyonunda değerlendirir."
    )
    add_bullet(doc, "Katı kısıtlar: yapı içinde kalmama, enkaz güvenlik çeperi, blokajsız besleme bacağı, en az bir açık hedef ve farklı sağlam cephelerde montaj.")
    add_bullet(doc, "Yumuşak ölçütler: toplam yol uzunluğu, açık kapsama oranı, minimum/ortalama aday kalitesi, uydu azimutuna erişim ve küme merkezine yakınlık.")
    add_bullet(doc, "Cephe adayları dış normal vektörü boyunca ötelenir; normal doğrultusu terminal, hedef küme ve sabit uydu azimutu ile açısal olarak karşılaştırılır.")
    add_bullet(doc, "Aday seti optimizasyonu tekil en yüksek puanı seçmek yerine mekânsal çeşitlilik, kapsama katkısı ve farklı bina zorunluluğunu birlikte gözetir.")
    add_callout(
        doc,
        "Kanal modeli tanımı",
        "İki-atlamalı röle yolu yalnızca bu aşamada açıkça Terminal→IRS ve IRS→depremzede bacakları "
        "olarak tanımlanır. Belgenin devamında aynı yapılar sırasıyla “besleme bacağı” ve "
        "“yansıtılmış erişim bacağı” olarak adlandırılmıştır.",
        fill=PALE,
        accent=BLUE,
    )

    new_page(doc)
    add_heading(doc, "5. IRS mühendislik metrikleri ve kanal uygunluğu", 1)
    add_screenshot(
        doc,
        "04-irs-ozet-penceresi.png",
        "Şekil 5. Seçilen yansıtıcı yüzey için kanal, kapsama, montaj ve kalite metrikleri.",
        width=6.25,
    )
    add_body(
        doc,
        "Ayrıntı penceresi, seçilen yansıtıcı yüzeyin özellik vektörünü operasyonel bir mühendislik "
        "özetine dönüştürür. Besleme bacağındaki LoS değeri ikili fiziksel uygunluğu; yansıtılmış erişim "
        "oranı ise kapsama yarıçapındaki hedeflerin ne kadarına obstrüksiyonsuz ulaşılabildiğini temsil eder."
    )
    add_bullet(doc, "Kanal durumu CLEAR, PARTIAL_NLoS ve FULL_NLoS sınıflarıyla raporlanır; tam blokaj bulunan aday uygun çözüm kümesinden çıkarılır.")
    add_bullet(doc, "Yansıma verimi, gelme ve çıkış açılarının kosinüs tabanlı bileşimiyle kestirilir; cephe normaline ters geometriler düşük verim üretir.")
    add_bullet(doc, "Kestirimsel link kazancı; IRS açıklık kazancı, minimum LoS faktörü ve toplam yol uzunluğuna bağlı sapma kaybının birleşimidir.")
    add_bullet(doc, "Bileşik kalite skoru; LoS, mesafe normu, montaj yüksekliği, cephe hizası, link bütçesi ve açık kapsama bileşenlerinin ağırlıklı toplamıdır.")
    add_bullet(doc, "Montaj azimutu, cephe yönü, tahmini yükseklik ve ev sahibi bina kimliği saha keşfi için izlenebilir meta-veri olarak saklanır.")
    add_callout(
        doc,
        "Model yorumu",
        "Arayüzdeki dB değeri tam uydu link bütçesi değildir. EIRP, G/T, alıcı gürültü sıcaklığı, "
        "polarizasyon kaybı, yağmur zayıflaması ve bant genişliği modele dahil edilmediğinden değer "
        "yalnızca adaylar arası göreli kanal uygunluğu için kullanılmalıdır.",
        fill="FFF8E8",
        accent=AMBER,
    )

    new_page(doc)
    add_heading(doc, "6. Görsel doğrulama ve karar otoritesi sınırı", 1)
    add_screenshot(
        doc,
        "06-gorsel-dogrulama.png",
        "Şekil 6. Render edilmiş saha çözümünün ikincil semantik tutarlılık denetimi.",
        width=6.25,
    )
    add_body(
        doc,
        "Görsel doğrulama katmanı, harita üzerinde beklenen nesne sınıflarının, işaretçi dağılımının ve "
        "bağlantı çizgilerinin genel tutarlılığını inceler. Bu işlem bir bilgisayarlı görü değerlendirmesidir; "
        "geometrik çözücüden gelen koordinatları, kanal sınıflarını veya uygunluk kararını değiştirmez."
    )
    add_bullet(doc, "Yerel motor: koordinat üretimi, kısıt kontrolü, amaç fonksiyonu, set seçimi ve nihai fiziksel uygunluk.")
    add_bullet(doc, "Açıklama katmanı: sayısal çıktının insan tarafından denetlenebilir doğal dil gerekçesine dönüştürülmesi.")
    add_bullet(doc, "Reranking katmanı: yalnızca eşik ve kısıt kontrollerini geçen çözümlerin operasyonel öncelik sırasının düzenlenmesi.")
    add_bullet(doc, "Görsel denetim: eksik işaretçi, belirgin çizgi tutarsızlığı veya sunum hatası gibi semantik anomalilerin bildirilmesi.")
    add_callout(
        doc,
        "Fail-safe ilke",
        "Yapay zekâ servisinin yanıt vermemesi, temel yerleşim sonucunu geçersiz kılmaz. Sistem yerel "
        "açıklama ve yerel sıralama ile çalışmaya devam eder; böylece haberleşme bağımlılığı kaynaklı "
        "tek hata noktası oluşturulmaz.",
    )

    new_page(doc)
    add_heading(doc, "7. ÖTR sistemiyle mühendislik bağlamı", 1)
    add_screenshot(
        doc,
        "07-proje-detaylari.png",
        "Şekil 7. Yazılım yerleşim motoru ile fiziksel hareketli uydu terminali alt sistemlerinin bağlamsal eşleştirilmesi.",
    )
    add_body(
        doc,
        "Proje Detayları ekranı, karar destek yazılımını 3GPP Release 17/18 NTN hedefleri ve fiziksel "
        "terminal mimarisiyle ilişkilendirir. ÖTR bağlamında platform; doğrudan tahrikli BLDC motorlar, "
        "manyetik enkoder, IMU tabanlı sensör füzyonu, Genişletilmiş Kalman Filtresi (EKF), kaskad "
        "PI/PID yönelim kontrolü ve slipring üzerinden sürekli azimut hareketi gibi alt sistemler içerir."
    )
    add_bullet(doc, "Mekanik yönelim katmanı, bozucu hareket altında ±8° bastırma ve 8 saniyenin altında yeniden kilitlenme hedefleriyle ilişkilidir.")
    add_bullet(doc, "Elektronik beamsteering katmanı, RIS/IRS hücrelerinin faz profilini değiştirerek yansıtılan dalga cephesini hedef doğrultuya yönlendirir.")
    add_bullet(doc, "Yerleşim motoru; mekanik kontrol döngüsünü simüle etmez, uygun kurulum koordinatı ve cephe geometrisi üreterek üst seviye görev planlaması sağlar.")
    add_bullet(doc, "20 kg / 140 W tasarım bütçesi fiziksel prototip kısıtıdır; yazılım skoru doğrudan kütle veya güç tüketimi hesabı içermez.")
    add_heading(doc, "Akademik sonuç", 2)
    add_body(
        doc,
        "ARES-Reflect, NLoS afet sahalarında akıllı yansıtıcı yüzey destekli erişim topolojisini "
        "oluşturmak için deterministik jeo-uzamsal analiz, kısıtlı kombinatoryal seçim ve açıklanabilir "
        "çok ölçütlü puanlamayı birleştiren yerel-öncelikli bir karar destek mimarisidir. Üretilen sonuçlar "
        "saha keşfi için aday konum niteliğindedir; nihai kabul için 3B ışın izleme, ayrıntılı RF link bütçesi, "
        "Fresnel açıklığı, yapısal montaj analizi ve prototip ölçümleri gereklidir."
    )


def finalize(doc):
    core = doc.core_properties
    core.title = "ARES-Reflect Arayüz ve Sistem Çalışma Akışı"
    core.subject = "Light mod arayüz ekran görüntüleriyle teknik sistem açıklaması"
    core.author = "ARES-Reflect Proje Ekibi"
    core.keywords = "ARES-Reflect, uydu terminali, IRS, RIS, afet haberleşmesi, arayüz"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    document = setup_document()
    add_academic_cover(document)
    add_academic_pages(document)
    finalize(document)
